"""将 PaperAgent 的结构化研究结果导出为可阅读的 Word/PDF 报告。"""

from __future__ import annotations

import re
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any, Iterable
from uuid import uuid4

BLUE, INK, MUTED = "3657DC", "263247", "68738A"


def _safe_name(value: str) -> str:
    name = re.sub(r"[^\w\-]+", "_", value.strip())[:50].strip("_") or "research_report"
    return f"{name}_{datetime.now():%Y%m%d_%H%M%S}_{uuid4().hex[:6]}"


def _clean_markdown(text: str) -> str:
    text = re.sub(r"\[([^]]+)]\((https?://[^)]+)\)", r"\1（\2）", text)
    return re.sub(r"[*_`]", "", text).strip()


def _lines(text: str) -> Iterable[tuple[str, str]]:
    """将常见 Markdown 行归一为标题、列表或正文。"""
    for raw in str(text).splitlines():
        line = raw.strip()
        if not line:
            continue
        if match := re.match(r"^(#{1,3})\s+(.+)$", line):
            yield f"heading_{min(len(match.group(1)) + 1, 3)}", _clean_markdown(match.group(2))
        elif match := re.match(r"^[-*+]\s+(.+)$", line):
            yield "bullet", _clean_markdown(match.group(1))
        elif match := re.match(r"^\d+[.)]\s+(.+)$", line):
            yield "number", _clean_markdown(match.group(1))
        else:
            yield "body", _clean_markdown(line)


def _summary_rows(payload: dict[str, Any]) -> list[tuple[str, str]]:
    metadata = payload.get("metadata") or {}
    verification = metadata.get("answer_verification") or metadata.get("citation_validation") or {}
    verification_status = verification.get("status") or ("通过" if verification.get("passed") else "未提供")
    return [
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("任务类型", str(payload.get("task_type") or "research")),
        ("检索范围", str(metadata.get("retrieval_source") or metadata.get("retrieval_scope") or "未提供")),
        ("证据数量", str(len(payload.get("papers") or []))),
        ("质量校验", str(verification_status)),
        ("Trace ID", str(payload.get("trace_id") or "未提供")),
    ]


def _set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    """固定 Word 表格的 DXA 总宽、列宽和缩进，避免跨渲染器漂移。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", sum(widths_dxa)), ("tblInd", 120)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def export_docx(payload: dict[str, Any], output_dir: str | Path) -> Path:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    path = output / f"{_safe_name(str(payload.get('title') or 'research_report'))}.docx"
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    tokens = {
        "Normal": (10.5, INK, 0, 6), "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6), "Heading 3": (12, "1F4D78", 8, 4),
        "List Bullet": (10.5, INK, 0, 6), "List Number": (10.5, INK, 0, 6),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        for font_key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
            style._element.rPr.rFonts.set(qn(font_key), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("PAPERAGENT · EVIDENCE-DRIVEN RESEARCH")
    run.bold, run.font.size = True, Pt(9)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run(str(payload.get("title") or "PaperAgent 研究报告"))
    run.bold, run.font.size, run.font.name = True, Pt(25), "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = RGBColor.from_string("111B34")
    subtitle = doc.add_paragraph("基于检索证据、引用校验与质量门控生成")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    rows = _summary_rows(payload)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    _set_table_geometry(table, [2160, 7200])
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text, row.cells[1].text = label, value
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].paragraphs[0].runs[0].bold = True
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F2F4F7")
        row.cells[0]._tc.get_or_add_tcPr().append(shading)

    doc.add_heading("研究问题", level=1)
    doc.add_paragraph(str(payload.get("query") or "未提供"))
    doc.add_heading("研究结论", level=1)
    for kind, text in _lines(str(payload.get("answer") or "")):
        if kind.startswith("heading_"):
            doc.add_heading(text, level=int(kind[-1]))
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(text, style="List Number")
        else:
            doc.add_paragraph(text)

    doc.add_heading("论文与证据", level=1)
    papers = payload.get("papers") or []
    if not papers:
        doc.add_paragraph("本次报告未附带可导出的论文证据。")
    for index, paper in enumerate(papers, 1):
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before, heading.paragraph_format.space_after = Pt(8), Pt(3)
        heading.add_run(f"{index}. {paper.get('title') or '未命名论文'}").bold = True
        authors = paper.get("authors") or []
        details = " · ".join(str(item) for item in (
            ", ".join(authors[:4]) if isinstance(authors, list) else authors,
            paper.get("source"), paper.get("year"), f"第 {paper.get('page')} 页" if paper.get("page") else "",
        ) if item)
        if details:
            detail = doc.add_paragraph(details)
            detail.paragraph_format.space_after = Pt(3)
            detail.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        if paper.get("content"):
            doc.add_paragraph(str(paper["content"])[:900])
        link = paper.get("pdf_url") or paper.get("entry_id")
        if link:
            doc.add_paragraph(f"来源：{link}")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("PaperAgent · 可验证科研工作流生成")
    doc.core_properties.author = "PaperAgent"
    doc.core_properties.title = str(payload.get("title") or "PaperAgent 研究报告")
    doc.core_properties.subject = "科研证据分析报告"
    doc.save(path)
    return path


def _pdf_text(value: Any) -> str:
    return escape(_clean_markdown(str(value))).replace("\n", "<br/>")


def _register_pdf_font() -> str:
    """优先嵌入本机中文 TrueType 字体，缺失时回退到标准 CID 字体。"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont

    for name, path in (("PaperAgentCN", Path("C:/Windows/Fonts/msyh.ttc")), ("PaperAgentCN", Path("C:/Windows/Fonts/simhei.ttf"))):
        if path.exists():
            try:
                pdfmetrics.registerFont(TTFont(name, str(path), subfontIndex=0))
                return name
            except Exception:
                continue
    try:
        pdfmetrics.getFont("STSong-Light")
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light"


def export_pdf(payload: dict[str, Any], output_dir: str | Path) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    path = output / f"{_safe_name(str(payload.get('title') or 'research_report'))}.pdf"
    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("CN", parent=styles["BodyText"], fontName=font_name, fontSize=10.5, leading=16, textColor=colors.HexColor(f"#{INK}"), spaceAfter=7)
    h1 = ParagraphStyle("CNH1", parent=body, fontSize=17, leading=23, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=12, spaceAfter=8)
    h2 = ParagraphStyle("CNH2", parent=body, fontSize=13, leading=18, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=9, spaceAfter=5)
    bullet = ParagraphStyle("BulletCN", parent=body, leftIndent=18, firstLineIndent=-9, bulletIndent=8)
    title_style = ParagraphStyle("TitleCN", parent=h1, fontSize=24, leading=31, alignment=TA_CENTER, textColor=colors.HexColor("#111B34"), spaceAfter=8)
    sub_style = ParagraphStyle("SubCN", parent=body, alignment=TA_CENTER, textColor=colors.HexColor(f"#{MUTED}"), spaceAfter=14)
    story = [
        Paragraph("PAPERAGENT · EVIDENCE-DRIVEN RESEARCH", ParagraphStyle("Kicker", parent=sub_style, fontSize=8.5, textColor=colors.HexColor(f"#{BLUE}"))),
        Paragraph(_pdf_text(payload.get("title") or "PaperAgent 研究报告"), title_style),
        Paragraph("基于检索证据、引用校验与质量门控生成", sub_style),
    ]
    summary = Table([[Paragraph(_pdf_text(k), body), Paragraph(_pdf_text(v), body)] for k, v in _summary_rows(payload)], colWidths=[100, 368])
    summary.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font_name), ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F4F7")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor(f"#{INK}")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), .4, colors.HexColor("#D5DAE5")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.extend([summary, Spacer(1, 12), Paragraph("研究问题", h1), Paragraph(_pdf_text(payload.get("query") or "未提供"), body), Paragraph("研究结论", h1)])
    for kind, text in _lines(str(payload.get("answer") or "")):
        if kind.startswith("heading_"):
            story.append(Paragraph(_pdf_text(text), h2))
        elif kind in {"bullet", "number"}:
            story.append(Paragraph(_pdf_text(text), bullet, bulletText="•" if kind == "bullet" else "-"))
        else:
            story.append(Paragraph(_pdf_text(text), body))

    papers = payload.get("papers") or []
    story.extend([PageBreak(), Paragraph("论文与证据", h1)])
    if not papers:
        story.append(Paragraph("本次报告未附带可导出的论文证据。", body))
    for index, paper in enumerate(papers, 1):
        story.append(Paragraph(f"{index}. {_pdf_text(paper.get('title') or '未命名论文')}", h2))
        authors = paper.get("authors") or []
        details = " · ".join(str(item) for item in (
            ", ".join(authors[:4]) if isinstance(authors, list) else authors,
            paper.get("source"), paper.get("year"), f"第 {paper.get('page')} 页" if paper.get("page") else "",
        ) if item)
        if details:
            story.append(Paragraph(_pdf_text(details), ParagraphStyle("DetailCN", parent=body, textColor=colors.HexColor(f"#{MUTED}"), fontSize=9.3)))
        if paper.get("content"):
            story.append(Paragraph(_pdf_text(str(paper["content"])[:900]), body))
        link = paper.get("pdf_url") or paper.get("entry_id")
        if link:
            story.append(Paragraph(f"来源：{_pdf_text(link)}", body))

    def footer(canvas, document):
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawCentredString(LETTER[0] / 2, 24, f"PaperAgent 研究报告 · 第 {document.page} 页")
        canvas.restoreState()

    SimpleDocTemplate(str(path), pagesize=LETTER, rightMargin=72, leftMargin=72, topMargin=64, bottomMargin=48, title=str(payload.get("title") or "PaperAgent 研究报告"), author="PaperAgent").build(story, onFirstPage=footer, onLaterPages=footer)
    return path
