from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader

from app.api import app
from reports.exporter import export_docx, export_pdf


def _payload():
    return {
        "title": "Agent 架构研究报告", "query": "比较 ReAct 与 Reflexion",
        "answer": "# 核心结论\nReAct 结合推理与行动。\n\n## 证据边界\n- 结论必须由论文支持。",
        "task_type": "compare", "trace_id": "trace-demo",
        "papers": [{"title": "ReAct", "source": "personal_library", "year": 2023,
                    "page": 2, "content": "Reasoning and acting are interleaved.",
                    "pdf_url": "https://arxiv.org/abs/2210.03629"}],
        "metadata": {"retrieval_source": "hybrid_personal_online"},
    }


def test_docx_report_contains_readable_sections_and_evidence(tmp_path):
    path = export_docx(_payload(), tmp_path)
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert path.suffix == ".docx"
    assert "核心结论" in text and "论文与证据" in text and "ReAct" in text
    assert len(document.tables) == 1
    assert len(document.tables[0].rows) == 6


def test_pdf_report_is_reopenable_and_multipage(tmp_path):
    path = export_pdf(_payload(), tmp_path)
    reader = PdfReader(path)
    assert path.suffix == ".pdf"
    assert len(reader.pages) >= 2
    assert path.stat().st_size > 2000


def test_report_export_api_returns_downloadable_docx_and_pdf(tmp_path, monkeypatch):
    from core.config import settings
    monkeypatch.setattr(settings, "REPORT_OUTPUT_DIR", str(tmp_path))
    client = TestClient(app)
    for report_format, media in (("docx", "application/vnd.openxmlformats"), ("pdf", "application/pdf")):
        response = client.post(f"/reports/export/{report_format}", json=_payload())
        assert response.status_code == 200
        assert media in response.headers["content-type"]
        assert response.content


def test_report_export_uses_unique_names_and_rejects_unknown_format(tmp_path):
    first = export_docx(_payload(), tmp_path)
    second = export_docx(_payload(), tmp_path)
    assert first.name != second.name
    response = TestClient(app).post("/reports/export/txt", json=_payload())
    assert response.status_code == 400
